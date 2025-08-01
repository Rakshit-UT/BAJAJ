"""
Document Processing Service
Handles PDF, DOCX, and Email parsing with clause-level extraction
"""

import fitz  # PyMuPDF
import requests
import tempfile
import os
from docx import Document
import email
from email.mime.text import MIMEText
import logging
from typing import List, Dict, Any, Optional
from urllib.parse import urlparse
import re
import spacy
from pathlib import Path

logger = logging.getLogger(__name__)

class DocumentChunk:
    """Represents a document chunk with metadata"""

    def __init__(self, text: str, chunk_id: str, source_meta: Dict[str, Any]):
        self.text = text
        self.chunk_id = chunk_id
        self.source_meta = source_meta

    def to_dict(self) -> Dict[str, Any]:
        return {
            "text": self.text,
            "chunk_id": self.chunk_id,
            "source_meta": self.source_meta
        }

class DocumentProcessor:
    """Main document processing service"""

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        self.chunk_size = 1000  # characters
        self.overlap = 200  # character overlap between chunks

        # Try to load spaCy model, fallback if not available
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy model not found, using basic text processing")
            self.nlp = None

    async def process_document_from_url(self, url: str) -> List[DocumentChunk]:
        """
        Download and process document from URL

        Args:
            url: Document URL (PDF, DOCX, etc.)

        Returns:
            List of DocumentChunk objects
        """
        logger.info(f"Processing document from URL: {url}")

        try:
            # Download the document
            response = requests.get(url, timeout=30)
            response.raise_for_status()

            # Determine file type from URL or content-type
            file_extension = self._get_file_extension(url, response.headers.get('content-type', ''))

            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                tmp_file.write(response.content)
                tmp_file_path = tmp_file.name

            try:
                # Process based on file type
                if file_extension == '.pdf':
                    chunks = self._process_pdf(tmp_file_path, url)
                elif file_extension in ['.docx', '.doc']:
                    chunks = self._process_docx(tmp_file_path, url)
                else:
                    chunks = self._process_text(tmp_file_path, url)

                logger.info(f"Successfully processed document into {len(chunks)} chunks")
                return chunks

            finally:
                # Clean up temporary file
                os.unlink(tmp_file_path)

        except Exception as e:
            logger.error(f"Error processing document from URL: {str(e)}")
            raise Exception(f"Failed to process document: {str(e)}")

    def _get_file_extension(self, url: str, content_type: str) -> str:
        """Determine file extension from URL or content type"""
        # First try URL
        parsed_url = urlparse(url)
        path = parsed_url.path.lower()

        for ext in self.supported_formats:
            if path.endswith(ext):
                return ext

        # Try content type
        content_type_map = {
            'application/pdf': '.pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'application/msword': '.doc',
            'text/plain': '.txt'
        }

        return content_type_map.get(content_type.lower(), '.pdf')  # Default to PDF

    def _process_pdf(self, file_path: str, source_url: str) -> List[DocumentChunk]:
        """Process PDF document"""
        chunks = []

        try:
            doc = fitz.open(file_path)

            for page_num in range(len(doc)):
                page = doc[page_num]

                # Extract text blocks with better structure preservation
                blocks = page.get_text("dict")
                page_text = ""

                for block in blocks.get("blocks", []):
                    if "lines" in block:  # Text block
                        for line in block["lines"]:
                            for span in line["spans"]:
                                page_text += span["text"] + " "
                        page_text += "\n"

                # Clean up text
                page_text = self._clean_text(page_text)

                if page_text.strip():
                    # Create chunks for this page
                    page_chunks = self._create_text_chunks(
                        page_text,
                        {
                            "source_url": source_url,
                            "page": page_num + 1,
                            "document_type": "pdf",
                            "total_pages": len(doc)
                        }
                    )
                    chunks.extend(page_chunks)

            doc.close()

        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise Exception(f"Failed to process PDF: {str(e)}")

        return chunks

    def _process_docx(self, file_path: str, source_url: str) -> List[DocumentChunk]:
        """Process DOCX document"""
        chunks = []

        try:
            doc = Document(file_path)
            full_text = ""

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n\n"

            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text += " | ".join(row_text) + "\n"

            # Clean up text
            full_text = self._clean_text(full_text)

            if full_text.strip():
                chunks = self._create_text_chunks(
                    full_text,
                    {
                        "source_url": source_url,
                        "document_type": "docx"
                    }
                )

        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise Exception(f"Failed to process DOCX: {str(e)}")

        return chunks

    def _process_text(self, file_path: str, source_url: str) -> List[DocumentChunk]:
        """Process plain text document"""
        chunks = []

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            text = self._clean_text(text)

            if text.strip():
                chunks = self._create_text_chunks(
                    text,
                    {
                        "source_url": source_url,
                        "document_type": "text"
                    }
                )

        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            raise Exception(f"Failed to process text file: {str(e)}")

        return chunks

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s.,;:!?()-]', ' ', text)

        # Normalize spacing around punctuation
        text = re.sub(r'\s*([.,;:!?])\s*', r'\1 ', text)

        return text.strip()

    def _create_text_chunks(self, text: str, base_metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Create overlapping text chunks"""
        chunks = []

        if len(text) <= self.chunk_size:
            # Single chunk
            chunk = DocumentChunk(
                text=text,
                chunk_id=f"chunk_0",
                source_meta={**base_metadata, "chunk_index": 0}
            )
            chunks.append(chunk)
        else:
            # Multiple chunks with overlap
            start = 0
            chunk_index = 0

            while start < len(text):
                end = start + self.chunk_size

                # If not the last chunk, try to break at word boundary
                if end < len(text):
                    # Find the last space within overlap region
                    break_point = text.rfind(' ', end - self.overlap, end)
                    if break_point > start:
                        end = break_point

                chunk_text = text[start:end].strip()

                if chunk_text:
                    chunk = DocumentChunk(
                        text=chunk_text,
                        chunk_id=f"chunk_{chunk_index}",
                        source_meta={
                            **base_metadata,
                            "chunk_index": chunk_index,
                            "start_char": start,
                            "end_char": end
                        }
                    )
                    chunks.append(chunk)
                    chunk_index += 1

                # Move start position
                start = end - self.overlap if end < len(text) else len(text)

        return chunks

    def extract_clauses(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """
        Extract and identify specific clause types from chunks
        Useful for insurance, legal, and compliance documents
        """
        clause_patterns = {
            "definitions": r"\b(definitions?|defined terms?)\b",
            "coverage": r"\b(coverage|covered|covers|benefits?)\b",
            "exclusions": r"\b(exclusions?|excluded|not covered)\b",
            "conditions": r"\b(conditions?|requirements?|eligibility)\b",
            "claims": r"\b(claims?|claim process|how to claim)\b",
            "premiums": r"\b(premiums?|payment|due date)\b",
            "waiting_period": r"\b(waiting period|waiting time)\b",
            "deductibles": r"\b(deductibles?|co-pay|copayment)\b"
        }

        enhanced_chunks = []

        for chunk in chunks:
            clause_types = []
            text_lower = chunk.text.lower()

            for clause_type, pattern in clause_patterns.items():
                if re.search(pattern, text_lower, re.IGNORECASE):
                    clause_types.append(clause_type)

            # Add clause type information to metadata
            enhanced_meta = chunk.source_meta.copy()
            enhanced_meta["clause_types"] = clause_types

            enhanced_chunk = DocumentChunk(
                text=chunk.text,
                chunk_id=chunk.chunk_id,
                source_meta=enhanced_meta
            )
            enhanced_chunks.append(enhanced_chunk)

        return enhanced_chunks

    def process_email(self, email_content: str) -> List[DocumentChunk]:
        """Process email content"""
        chunks = []

        try:
            msg = email.message_from_string(email_content)

            # Extract headers
            headers = {
                "from": msg.get("From", ""),
                "to": msg.get("To", ""),
                "subject": msg.get("Subject", ""),
                "date": msg.get("Date", "")
            }

            # Extract body
            body_text = ""
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            body_text += payload.decode('utf-8', errors='ignore')
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    body_text = payload.decode('utf-8', errors='ignore')

            # Clean text
            body_text = self._clean_text(body_text)

            if body_text.strip():
                chunks = self._create_text_chunks(
                    body_text,
                    {
                        "document_type": "email",
                        "headers": headers
                    }
                )

        except Exception as e:
            logger.error(f"Error processing email: {str(e)}")
            raise Exception(f"Failed to process email: {str(e)}")

        return chunks
